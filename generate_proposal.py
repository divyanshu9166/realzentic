from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Colour palette ────────────────────────────────────────────────────────────
BROWN_DARK   = RGBColor(0x4A, 0x2C, 0x2A)   # deep mahogany
BROWN_MID    = RGBColor(0x7B, 0x4F, 0x2E)   # warm walnut
GOLD         = RGBColor(0xC9, 0x8B, 0x2B)   # brass / gold
CREAM        = RGBColor(0xFD, 0xF6, 0xEC)   # off-white
DARK_TEXT    = RGBColor(0x1E, 0x1E, 0x1E)   # near-black
GRAY_MID     = RGBColor(0x55, 0x55, 0x55)   # body grey
GREEN_CHECK  = RGBColor(0x2E, 0x7D, 0x32)   # tick colour
RED_PAIN     = RGBColor(0xC6, 0x28, 0x28)   # pain point red

# ── Helper utilities ──────────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_cell_border(cell, border_color='C98B2B', side='bottom', width='12'):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    b = OxmlElement(f'w:{side}')
    b.set(qn('w:val'),   'single')
    b.set(qn('w:sz'),    width)
    b.set(qn('w:space'), '0')
    b.set(qn('w:color'), border_color)
    tcBorders.append(b)
    tcPr.append(tcBorders)

def paragraph_space(doc, before=0, after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    return p

def add_heading(doc, text, level=1, color=BROWN_DARK, size=22, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
    p   = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def add_subheading(doc, text, color=BROWN_MID, size=13):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def add_body(doc, text, color=DARK_TEXT, size=10.5, italic=False):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(size)
    run.font.color.rgb = color
    run.italic = italic
    return p

def add_bullet(doc, text, color=DARK_TEXT, bullet_char='✓', bullet_color=GREEN_CHECK, size=10.5):
    p   = doc.add_paragraph()
    p.paragraph_format.left_indent   = Inches(0.25)
    p.paragraph_format.space_before  = Pt(2)
    p.paragraph_format.space_after   = Pt(3)
    br = p.add_run(bullet_char + '  ')
    br.font.color.rgb = bullet_color
    br.font.size      = Pt(size)
    br.bold           = True
    tr = p.add_run(text)
    tr.font.color.rgb = color
    tr.font.size      = Pt(size)
    return p

def add_pain_bullet(doc, text):
    return add_bullet(doc, text, color=DARK_TEXT, bullet_char='✗', bullet_color=RED_PAIN)

def add_divider(doc, color_hex='C98B2B'):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def banner_table(doc, text, bg_hex='4A2C2A', fg=RGBColor(0xFF,0xFF,0xFF), size=20):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, bg_hex)
    cell.width = Inches(7)
    p   = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    run = p.add_run(text)
    run.bold           = True
    run.font.size      = Pt(size)
    run.font.color.rgb = fg
    doc.add_paragraph()
    return tbl

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
paragraph_space(doc, after=30)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('FurnitureCRM')
r.bold = True
r.font.size = Pt(42)
r.font.color.rgb = BROWN_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Your Complete Business Operating System')
r.font.size = Pt(18)
r.font.color.rgb = GOLD
r.bold = True

paragraph_space(doc, after=8)
add_divider(doc, 'C98B2B')
paragraph_space(doc, after=8)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('A Growth & Operations Platform Built Exclusively for Furniture Businesses')
r.font.size = Pt(13)
r.font.color.rgb = GRAY_MID
r.italic = True

paragraph_space(doc, after=40)

# Feature snapshot cards in a 3-col table
tbl = doc.add_table(rows=2, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
icons = [
    ('📈', 'Marketing\nAutomation',  'F5EFE6'),
    ('⚙️', 'Operations\nManagement', 'F5EFE6'),
    ('🤖', 'AI Voice\nAgent',         'F5EFE6'),
    ('💬', 'Omnichannel\nInbox',      'F5EFE6'),
    ('📦', 'Inventory\n& Billing',    'F5EFE6'),
    ('⭐', 'Reviews\n& Reputation',   'F5EFE6'),
]
for i, (icon, label, bg) in enumerate(icons):
    row_idx = i // 3
    col_idx = i % 3
    cell = tbl.cell(row_idx, col_idx)
    set_cell_bg(cell, bg)
    add_cell_border(cell, 'C98B2B', 'bottom', '8')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(icon + '\n')
    r.font.size = Pt(22)
    r2 = p.add_run(label)
    r2.font.size = Pt(10)
    r2.bold = True
    r2.font.color.rgb = BROWN_DARK

paragraph_space(doc, after=30)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Prepared exclusively for your furniture business  |  Confidential')
r.font.size = Pt(9)
r.font.color.rgb = GRAY_MID
r.italic = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('© 2025 FurnitureCRM  |  India')
r.font.size = Pt(9)
r.font.color.rgb = GRAY_MID

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
banner_table(doc, '01  |  Executive Summary')

add_body(doc,
    'Running a furniture store in India today means managing walk-in customers, WhatsApp enquiries, '
    'Instagram leads, custom orders, delivery schedules, GST billing, staff attendance, and supplier '
    'follow-ups — often with sticky notes, WhatsApp groups, and disconnected spreadsheets. '
    'The result: leads fall through, cash flow is unclear, and owners are buried in daily firefighting.',
    size=11)

add_body(doc,
    'FurnitureCRM was built to eliminate that chaos.  It is a single, purpose-built platform that '
    'connects every part of your furniture business — from the first customer enquiry to final payment '
    'collection — in one easy-to-use system.  With a built-in Indian-language Voice AI Agent for '
    'inbound and outbound calls, automated follow-ups, and real-time business dashboards, '
    'FurnitureCRM gives you full control without adding to your workload.',
    size=11)

paragraph_space(doc, after=6)

# 3 outcome boxes
tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
outcomes = [
    ('↑ 35%',  'More Leads\nConverted',   '4A2C2A', 'FFFFFF'),
    ('↓ 60%',  'Less Daily\nManual Work',  'C98B2B', 'FFFFFF'),
    ('2×',     'Faster Follow-up\nSpeed',  '7B4F2E', 'FFFFFF'),
]
for col_idx, (stat, label, bg, fg) in enumerate(outcomes):
    cell = tbl.cell(0, col_idx)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(12)
    r = p.add_run(stat + '\n')
    r.font.size = Pt(28)
    r.bold = True
    r.font.color.rgb = GOLD if bg == '4A2C2A' else RGBColor(0xFF,0xFF,0xFF)
    r2 = p.add_run(label)
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

paragraph_space(doc, after=10)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — PAIN POINTS
# ══════════════════════════════════════════════════════════════════════════════
banner_table(doc, '02  |  The Real Problems Furniture Owners Face Every Day', bg_hex='7B4F2E')

add_body(doc,
    'Before we show you what FurnitureCRM does, let us acknowledge the challenges you already live with. '
    'These are not hypothetical — they are the daily reality reported by furniture store owners across India.',
    size=10.5, italic=True, color=GRAY_MID)

# Pain points in 2-column table
pain_categories = [
    ('📉  Lead & Sales Leakage', [
        'Walk-in customers leave without being followed up — no record kept',
        'WhatsApp enquiries get buried in personal chats and are forgotten',
        'Sales staff verbally promise follow-ups that never happen',
        'No way to know which Instagram or Facebook ad is generating real buyers',
        'Quotations shared on WhatsApp — no tracking, no reminders, no closure',
    ]),
    ('⏰  Daily Management Overload', [
        'Owner must personally chase staff for daily sales updates',
        'Attendance tracked on registers or WhatsApp — easily manipulated',
        'No visibility on which staff member is performing vs. underperforming',
        'Custom order details scattered across WhatsApp threads and notebooks',
        'Multiple tools (Excel, Tally, WhatsApp) with no single source of truth',
    ]),
    ('💸  Billing & Cash Flow Blind Spots', [
        'Pending payments forgotten until months later',
        'Manual invoice creation — errors, delays, and GST miscalculations',
        'No alert when a high-value customer has an overdue balance',
        'Partial payments untracked — customer calls demanding proof of payment',
        'No visibility on daily/weekly/monthly revenue at a glance',
    ]),
    ('📦  Inventory & Operations Chaos', [
        'Products sold that are out of stock — customer dissatisfaction',
        'No reorder alerts — procurement handled on memory and gut feel',
        'Custom order measurements stored in physical notebooks — easily lost',
        'Delivery schedules managed via WhatsApp groups — constant confusion',
        'Production status unknown unless owner calls the factory directly',
    ]),
    ('📱  Customer Communication Breakdown', [
        'Enquiries coming from WhatsApp, Instagram DM, email, and website — no single inbox',
        'Slow response times drive potential buyers to competitors',
        'No post-sale follow-up — customers never asked for referrals or reviews',
        'Negative Google reviews go unnoticed for weeks',
        'Festival/season marketing campaigns done manually with no tracking',
    ]),
    ('📞  Phone Call Inefficiency', [
        'Staff waste hours on repetitive calls — price, availability, store hours',
        'Missed calls from hot leads — no callback mechanism',
        'No record of call conversations — disputes with customers',
        'Outbound follow-up calls skipped when staff are busy on the floor',
        'Language barrier — difficulty communicating in regional languages',
    ]),
]

for title, bullets in pain_categories:
    add_subheading(doc, title, color=BROWN_DARK, size=12)
    for b in bullets:
        add_pain_bullet(doc, b)
    paragraph_space(doc, after=4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — SOLUTION OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
banner_table(doc, '03  |  How FurnitureCRM Solves Every One of These Problems')

add_body(doc,
    'FurnitureCRM is not a generic CRM adapted for furniture.  It was designed ground-up for '
    'the Indian furniture retail and manufacturing context — with modules that match how you '
    'actually run your business.',
    size=11)

# Solution map table (problem → solution)
add_subheading(doc, 'Problem-to-Solution Map', color=BROWN_MID)

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
hdr_cells = tbl.rows[0].cells
for cell, text, bg in [(hdr_cells[0], 'Your Problem', 'C62828'), (hdr_cells[1], 'FurnitureCRM Solution', '2E7D32')]:
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

rows_data = [
    ('Leads lost from walk-ins',             'Walk-in module captures every visitor with instant follow-up assignment'),
    ('WhatsApp leads forgotten',             'Omnichannel Conversations inbox consolidates all channels in one view'),
    ('No follow-up accountability',          'Lead pipeline with stage tracking + automated follow-up reminders'),
    ('Cannot measure ad ROI',               'Source attribution on every lead — track which channel converts best'),
    ('Staff attendance manipulation',        'Digital clock-in/clock-out with worked-hour calculations'),
    ('Custom order chaos',                   'Dedicated Custom Orders module with measurements, timeline & payments'),
    ('Pending payment blind spots',         'Billing module with overdue alerts and pending collection dashboard'),
    ('GST billing errors',                   'GST-ready POS with auto tax calculations and invoice generation'),
    ('Inventory stockouts',                  'SKU-level inventory with low-stock and reorder alerts'),
    ('Slow phone response',                  'Indian Voice AI Agent handles inbound calls 24/7 in regional languages'),
    ('Missed outbound follow-ups',           'AI Agent makes outbound calls automatically at scheduled intervals'),
    ('Negative reviews unnoticed',           'Reviews module with low-rating escalation alerts and reply tracking'),
    ('No marketing measurement',             'Campaign module links spend to leads, conversions, and revenue'),
    ('Owner stuck in daily operations',      'Dashboard command center — full visibility in 60 seconds every morning'),
]
for prob, sol in rows_data:
    row = tbl.add_row()
    c0, c1 = row.cells[0], row.cells[1]
    set_cell_bg(c0, 'FFF8F8')
    set_cell_bg(c1, 'F0FFF0')
    p0 = c0.paragraphs[0]
    r0 = p0.add_run('✗  ' + prob)
    r0.font.size = Pt(9.5)
    r0.font.color.rgb = RED_PAIN
    p1 = c1.paragraphs[0]
    r1 = p1.add_run('✓  ' + sol)
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = GREEN_CHECK

paragraph_space(doc, after=10)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — VOICE AI AGENT
# ══════════════════════════════════════════════════════════════════════════════
banner_table(doc, '04  |  Indian Voice AI Agent — Your 24/7 Phone Representative', bg_hex='2E4A7B')

add_body(doc,
    'The FurnitureCRM Voice AI Agent is a trained, conversational AI that speaks naturally in '
    'Hindi, English, and major regional Indian languages.  It handles both inbound customer calls '
    'and outbound follow-up calls — so no lead is ever missed because a staff member was busy.',
    size=11)

add_subheading(doc, 'Inbound Call Capabilities', color=RGBColor(0x1A, 0x3A, 0x6E))
inbound = [
    'Answers calls 24/7 — even after store hours and on Sundays',
    'Responds to price enquiries, product availability, and store location questions',
    'Books showroom appointments and adds them directly to your CRM calendar',
    'Captures lead details (name, number, requirement) and creates a CRM lead automatically',
    'Transfers to a human agent if the query is complex or the customer requests it',
    'Handles calls in Hindi, English, Gujarati, Marathi, Tamil, Telugu, Kannada, and more',
    'Reduces staff phone burden by up to 70% for repetitive enquiry calls',
]
for item in inbound:
    add_bullet(doc, item, bullet_color=RGBColor(0x1A, 0x3A, 0x6E), bullet_char='▶')

add_subheading(doc, 'Outbound Call Capabilities', color=RGBColor(0x1A, 0x3A, 0x6E))
outbound = [
    'Automatically calls leads who have not been followed up within your defined window',
    'Reminds customers about pending quotation decisions with a polite follow-up call',
    'Calls customers whose orders are ready for pickup or delivery',
    'Sends payment reminder calls to customers with overdue balances',
    'Follows up post-delivery to collect reviews and satisfaction feedback',
    'All call outcomes are logged in the CRM Call Center module automatically',
]
for item in outbound:
    add_bullet(doc, item, bullet_color=RGBColor(0x1A, 0x3A, 0x6E), bullet_char='▶')

add_subheading(doc, 'Why This Matters for Your Business', color=RGBColor(0x1A, 0x3A, 0x6E))
add_body(doc,
    'A typical furniture store misses 3–5 inbound calls every day.  Each missed call is a potential '
    '₹15,000–₹2,00,000 sale gone to a competitor.  The Voice AI Agent ensures every call is answered, '
    'every follow-up is made, and every customer interaction is recorded — without hiring additional staff.',
    size=10.5)

# Stats box
tbl = doc.add_table(rows=1, cols=4)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
stats = [
    ('24/7',  'Call\nAvailability', '1A3A6E'),
    ('0',     'Missed\nCalls',      '2E4A7B'),
    ('100%',  'Calls\nLogged',      '3A5F9E'),
    ('8+',    'Indian\nLanguages',  '4A70B5'),
]
for col_idx, (stat, label, bg) in enumerate(stats):
    cell = tbl.cell(0, col_idx)
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    r = p.add_run(stat + '\n')
    r.font.size = Pt(24)
    r.bold = True
    r.font.color.rgb = GOLD
    r2 = p.add_run(label)
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)

paragraph_space(doc, after=10)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — FEATURE MODULES
# ══════════════════════════════════════════════════════════════════════════════
banner_table(doc, '05  |  Complete Feature Modules')

modules = [
    ('📊  Dashboard — Your Daily Command Center', [
        'Real-time KPI cards: revenue, conversion rate, average order value, pending collections',
        'Sales funnel view — see exactly where leads are dropping off',
        'Channel performance by lead source (WhatsApp, walk-in, Instagram, website)',
        'Action center: pending follow-ups + overdue invoices in one glance',
        'Best-selling products and live stock alert visibility',
        'One-minute morning review replaces 30-minute team WhatsApp updates',
    ]),
    ('🚶  Walk-in Capture — Stop Losing Showroom Visitors', [
        'Register every walk-in with name, contact, and interest in under 30 seconds',
        'Assign walk-in to a sales staff member instantly',
        'Status progression: Browsing → Interested → Quotation → Converted / Left',
        'Automatic follow-up scheduling if walk-in leaves without buying',
        'Measure actual showroom conversion rate — not just footfall',
    ]),
    ('🎯  Lead Management — Full Pipeline Visibility', [
        'Pipeline stages: New → Contacted → Showroom Visit → Quotation → Converted → Lost',
        'Source attribution — know which channel each lead came from',
        'Follow-up scheduling with due-date reminders and alerts',
        'Operational notes on every lead — full history at a glance',
        'Stage-level conversion reports to identify weak handoff points',
        'Accountability: every lead has an assigned owner and a next action',
    ]),
    ('👨‍💼  Staff Management — Know Your Team\'s Performance', [
        'Complete team directory with roles, status, and login credentials',
        'Digital attendance clock-in/clock-out with exact worked-hour calculations',
        'Team-level productivity and sales activity analytics',
        'Role-based access: Admin, Manager, and Staff levels',
        'Staff portal for self-service clock-in, task updates, and personal dashboards',
        'Connects individual effort to business outcomes — no more guesswork',
    ]),
    ('📅  Appointments — Never Miss a Scheduled Visit', [
        'Create, view, and manage all customer appointments in one calendar',
        'Status tracking: Scheduled → Confirmed → Completed → Cancelled',
        'Automated reminders to reduce no-shows',
        'Linked to lead records — full context before every meeting',
    ]),
    ('📦  Inventory Management — No More Stockouts', [
        'SKU-level product catalog with real-time stock quantities',
        'Low-stock alerts when items fall below your reorder threshold',
        'Inventory visibility for sales staff — no more promising unavailable products',
        'Category-wise organization and fast product search',
        'Integrated with billing and orders for automatic stock deduction',
    ]),
    ('🛒  Order Management — Full Order Lifecycle Control', [
        'Unified order management with customizable status transitions',
        'Payment status tracking per order — partial, paid, pending',
        'Source-level visibility — marketplace vs direct vs referral',
        'Fulfillment discipline with clear delivery and completion tracking',
        'Channel profitability analysis — know which source is most valuable',
    ]),
    ('🧾  Billing & POS — Fast, Accurate, GST-Ready', [
        'Professional invoice generation with your store branding',
        'POS checkout with discount, tax (GST), and total calculations',
        'Partial payment support with outstanding balance tracking',
        'Overdue payment alerts with escalation reminders',
        'Drafts module to save mid-transaction states — no data loss',
        'Speeds billing operations at busy counter periods',
    ]),
    ('🛠️  Custom Orders — Manage Bespoke Projects End-to-End', [
        'Custom order status timeline: Enquiry → Measurement → Production → Ready → Delivered',
        'Field-visit support with site measurement capture and photo notes',
        'Advance payment and installment tracking for large projects',
        'Assigned team member with full project history',
        'Production notes visible to both office and factory teams',
        'Reduces delays caused by missing site details or unclear specifications',
    ]),
    ('📧  Email Marketing — Re-engage and Retain Customers', [
        'Campaign builder with audience segmentation (past buyers, enquiries, etc.)',
        'One-click campaign sending with open/click tracking',
        'Festival and season offer campaigns — Diwali, New Year, Independence Day',
        'Dormant customer re-engagement — bring back customers who haven\'t bought in 6 months',
        'Post-purchase nurture sequences to encourage repeat business',
    ]),
    ('💬  Conversations — One Inbox for All Channels', [
        'Consolidates messages from WhatsApp, Instagram DM, website chat, and email',
        'AI-handled, Needs Human, and Resolved status tracking',
        'Unread counter and priority alerts in top navigation bar',
        'Outbound messaging directly from the CRM',
        'Full conversation history per customer — no context switching',
        'Faster response = higher conversion: industry data shows 78% of buyers choose the first responder',
    ]),
    ('📞  Call Center — Voice-Led Sales Control', [
        'Log all inbound and outbound calls with outcome notes',
        'Contact management with call history and follow-up context',
        'Integration with the Voice AI Agent for automated call logging',
        'Call outcome reports for sales performance measurement',
        'Compliance tracking — ensure every lead receives a call within defined SLA',
    ]),
    ('⭐  Reviews & Reputation — Protect Your Brand Online', [
        'Collect and display customer reviews from Google and other platforms',
        'Low-rating escalation: immediate alert when a 1 or 2-star review is posted',
        'Reply tracking — ensure no negative review goes unaddressed',
        'Review request campaigns to boost your Google rating proactively',
        'Local SEO impact: higher ratings = more organic discovery on Google Maps',
    ]),
    ('🤖  AI Recommend — Sell Smarter, Upsell More', [
        'AI-assisted furniture and room recommendation flows',
        'Helps staff suggest matching products for higher basket value',
        'Better product visualization support for customer decision-making',
        'Upsell companion products: sofas + coffee tables, beds + wardrobes',
    ]),
]

for title, points in modules:
    add_subheading(doc, title, color=BROWN_DARK, size=12)
    for pt in points:
        add_bullet(doc, pt)
    paragraph_space(doc, after=4)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — MARKETING TRANSFORMATION
# ══════════════════════════════════════════════════════════════════════════════
banner_table(doc, '06  |  Marketing Transformation — Know What Works, Stop What Doesn\'t', bg_hex='2E5C3A')

add_body(doc,
    'Most furniture stores spend money on Instagram ads, Google campaigns, and hoardings — '
    'but have no way to know which spend is generating actual sales.  FurnitureCRM changes that.',
    size=11)

add_subheading(doc, 'Lead Source Attribution', color=RGBColor(0x1A, 0x4A, 0x28))
attribution_points = [
    'Every lead is tagged with its source at the point of capture: Walk-in, WhatsApp, Instagram, Facebook, Google, Website, Referral, or Campaign',
    'Source-level conversion funnel — see which channel brings the most buyers, not just enquiries',
    'Cost-per-conversion tracking when linked to campaign spend data',
    'Identify your top-performing channel and double down; eliminate underperforming spend',
    'Weekly source performance reports for owner-level marketing decisions',
]
for pt in attribution_points:
    add_bullet(doc, pt, bullet_color=RGBColor(0x1A, 0x4A, 0x28), bullet_char='▶')

add_subheading(doc, 'Campaign Management', color=RGBColor(0x1A, 0x4A, 0x28))
campaign_points = [
    'Plan and record all marketing campaigns — digital and offline — in one place',
    'Track enquiries, visits, and conversions generated by each campaign',
    'Email marketing campaigns with open rates and click tracking',
    'Festival campaign calendar: Diwali, Navratri, New Year, Wedding Season',
    'Post-campaign reports to measure ROI before approving the next spend',
]
for pt in campaign_points:
    add_bullet(doc, pt, bullet_color=RGBColor(0x1A, 0x4A, 0x28), bullet_char='▶')

add_subheading(doc, 'Review & Reputation Marketing', color=RGBColor(0x1A, 0x4A, 0x28))
review_points = [
    'Systematic review collection — every delivered customer gets a review request',
    'Higher Google ratings drive more organic discovery on Google Maps — zero ad spend',
    'Negative review alert within minutes — respond before the damage spreads',
    'Review reply tracking — customers who receive responses often revise negative ratings',
    'Showcase top reviews in your email campaigns and WhatsApp marketing',
]
for pt in review_points:
    add_bullet(doc, pt, bullet_color=RGBColor(0x1A, 0x4A, 0x28), bullet_char='▶')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — AUTOMATION
# ══════════════════════════════════════════════════════════════════════════════
banner_table(doc, '07  |  Automation — Let the System Do the Repetitive Work', bg_hex='5C3A7B')

add_body(doc,
    'The biggest hidden cost in a furniture business is time — owner time, manager time, '
    'and staff time — spent on repetitive tasks that a system can handle automatically.  '
    'Here is exactly what FurnitureCRM automates for you:',
    size=11)

automation_sections = [
    ('Sales Automation', [
        'Auto-assign incoming leads to available staff based on round-robin or rules',
        'Follow-up due-date reminders: alert staff when a follow-up is overdue',
        'Lead stage escalation alerts: notify manager if a lead is stuck in one stage too long',
        'Conversation status management: auto-flag unanswered messages for human review',
        'Quotation expiry reminders: customer gets a gentle follow-up before the quote expires',
    ]),
    ('Operations Automation', [
        'Attendance logging: automatic worked-hour calculation from clock-in/clock-out data',
        'Stock alert: automatic low-stock notification when inventory drops below threshold',
        'Custom order timeline: move through production stages with status update notifications',
        'Appointment reminders: customer confirmation and staff notification before visits',
        'Daily action center: surfaces the top 5 urgent items every morning automatically',
    ]),
    ('Finance Automation', [
        'GST and tax calculations: automatic in POS — no manual computation',
        'Overdue invoice alerts: flag and escalate unpaid invoices after defined days',
        'Partial payment tracking: automatic outstanding balance calculation',
        'Pending collections dashboard: ranked by amount and age automatically',
        'Revenue KPIs: daily/weekly/monthly revenue updated in real-time',
    ]),
    ('Communication Automation (Voice AI)', [
        'Inbound call answering: Voice AI responds immediately, 24/7, in regional languages',
        'Lead capture from calls: AI extracts name, number, and requirement and creates CRM lead',
        'Appointment booking via call: customer books a showroom visit through the AI',
        'Outbound follow-up calls: AI calls leads on your behalf at scheduled intervals',
        'Payment reminder calls: polite automated reminder to customers with overdue balances',
        'Post-delivery feedback calls: collect reviews automatically after every delivery',
    ]),
]

for title, points in automation_sections:
    add_subheading(doc, title, color=RGBColor(0x3A, 0x1A, 0x5C))
    for pt in points:
        add_bullet(doc, pt, bullet_color=RGBColor(0x5C, 0x3A, 0x7B), bullet_char='⚡')

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — DAILY OPERATING RHYTHM
# ══════════════════════════════════════════════════════════════════════════════
banner_table(doc, '08  |  Your New Daily Operating Rhythm', bg_hex='4A2C2A')

add_body(doc,
    'Here is what your day looks like after implementing FurnitureCRM:',
    size=11, color=GRAY_MID, italic=True)

rhythm = [
    ('⏰  Morning (5 minutes)', [
        'Open Dashboard → review overnight enquiries, Voice AI call logs, and pending follow-ups',
        'Check action center → 3 overdue invoices, 2 follow-up calls due today, 1 stock alert',
        'Review today\'s appointments → 4 showroom visits scheduled',
        'Done. Team briefed. Priorities clear. No WhatsApp group needed.',
    ]),
    ('☀️  Day (Real-time)', [
        'Voice AI answers all inbound calls — you focus on floor selling',
        'Walk-in registrations captured by staff in under 30 seconds',
        'Conversations inbox handles WhatsApp and Instagram enquiries — AI responds first',
        'Billing done at POS in 2 minutes — GST invoice sent to customer\'s WhatsApp',
        'Custom order field team updates production status from the staff portal',
    ]),
    ('🌙  Evening (10 minutes)', [
        'Review daily revenue, new leads, and conversion update on dashboard',
        'Check unresolved conversations — respond to any that need human touch',
        'Confirm tomorrow\'s appointments — no-show risk managed',
        'Voice AI schedules outbound follow-up calls for tomorrow automatically',
    ]),
    ('📅  Weekly (30 minutes)', [
        'Funnel movement: how many leads converted vs lost this week',
        'Channel performance: which source brought the best buyers',
        'Staff productivity and attendance review',
        'Stock risk check: what needs to be reordered',
        'Campaign performance: was last week\'s promotion worth the spend',
    ]),
]

for title, points in rhythm:
    add_subheading(doc, title, color=BROWN_DARK)
    for pt in points:
        add_bullet(doc, pt)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — BUSINESS IMPACT
# ══════════════════════════════════════════════════════════════════════════════
banner_table(doc, '09  |  Expected Business Impact')

add_body(doc,
    'Based on how furniture businesses use CRM platforms in the Indian market, here are the '
    'tangible outcomes you can expect within 60–90 days of full adoption:',
    size=11)

impact_table_data = [
    ('Area',              'Before FurnitureCRM',            'After FurnitureCRM',              'Impact'),
    ('Lead Follow-up',   'Ad-hoc, verbal, forgotten',       'Automated reminders, 100% tracked','↑ 30–40% conversion'),
    ('Missed Calls',     '3–5 per day, no callback',        'Voice AI answers all, 24/7',       '₹0 revenue leakage'),
    ('Billing Speed',    '10–15 min per invoice',           '2 min POS checkout',               '↓ 80% billing time'),
    ('Collection',       'Chasing customers manually',      'Automated overdue alerts + AI calls','↑ 25% collection rate'),
    ('Staff Mgmt',       'Attendance on paper/WhatsApp',    'Digital with worked-hour accuracy', '↓ 90% admin effort'),
    ('Marketing ROI',    'Unknown which ads work',          'Source attribution per lead',       'Smarter spend decisions'),
    ('Google Reviews',   'Reactive to negatives',           'Systematic collection + monitoring','↑ 0.5–1 star improvement'),
    ('Owner Daily Time', '3–4 hrs in operations/WhatsApp',  '30 min dashboard review',          'Focus on growth, not ops'),
]

tbl = doc.add_table(rows=len(impact_table_data), cols=4)
tbl.style = 'Table Grid'
for row_idx, row_data in enumerate(impact_table_data):
    for col_idx, cell_text in enumerate(row_data):
        cell = tbl.cell(row_idx, col_idx)
        if row_idx == 0:
            set_cell_bg(cell, '4A2C2A')
            r = cell.paragraphs[0].add_run(cell_text)
            r.bold = True
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        else:
            bg = 'FDF6EC' if row_idx % 2 == 0 else 'FFFFFF'
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            r = p.add_run(cell_text)
            r.font.size = Pt(9.5)
            if col_idx == 3:
                r.font.color.rgb = GREEN_CHECK
                r.bold = True
            else:
                r.font.color.rgb = DARK_TEXT

paragraph_space(doc, after=10)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — WHY FURNITURECRM
# ══════════════════════════════════════════════════════════════════════════════
banner_table(doc, '10  |  Why FurnitureCRM Over Generic Tools', bg_hex='7B4F2E')

comparison = [
    ('Feature',                      'Spreadsheets / WhatsApp', 'Generic CRM\n(Zoho, Salesforce)', 'FurnitureCRM'),
    ('Built for furniture business',  '✗',  '✗ (adapted)',  '✓ Purpose-built'),
    ('Walk-in capture module',        '✗',  '✗',            '✓'),
    ('Custom order project tracking', '✗',  '✗',            '✓'),
    ('Indian Voice AI Agent',         '✗',  '✗',            '✓'),
    ('Regional language support',     '✗',  'Partial',      '✓ 8+ languages'),
    ('GST-ready billing & POS',       'Manual', 'Add-on cost', '✓ Built-in'),
    ('Staff attendance module',       '✗',  '✗',            '✓'),
    ('Furniture AI recommendations',  '✗',  '✗',            '✓'),
    ('Setup complexity',              'None (but useless)', 'High (3–6 months)', 'Low (1–2 days)'),
    ('Cost',                          'Free but costly in lost revenue', 'High (per-user fees)', 'Affordable fixed pricing'),
    ('Support',                       'None', 'Ticket-based, slow', 'Dedicated onboarding'),
]

tbl = doc.add_table(rows=len(comparison), cols=4)
tbl.style = 'Table Grid'
for row_idx, row_data in enumerate(comparison):
    for col_idx, cell_text in enumerate(row_data):
        cell = tbl.cell(row_idx, col_idx)
        if row_idx == 0:
            bg = '4A2C2A'
            set_cell_bg(cell, bg)
            r = cell.paragraphs[0].add_run(cell_text)
            r.bold = True; r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        else:
            if col_idx == 3:
                set_cell_bg(cell, 'F0FFF0')
                r = cell.paragraphs[0].add_run(cell_text)
                r.font.size = Pt(9.5)
                r.font.color.rgb = GREEN_CHECK
                r.bold = True
            elif col_idx == 0:
                set_cell_bg(cell, 'FDF6EC')
                r = cell.paragraphs[0].add_run(cell_text)
                r.font.size = Pt(9.5)
                r.bold = True
                r.font.color.rgb = BROWN_DARK
            else:
                bg = 'FFF8F8' if row_idx % 2 == 0 else 'FFFFFF'
                set_cell_bg(cell, bg)
                r = cell.paragraphs[0].add_run(cell_text)
                r.font.size = Pt(9.5)
                r.font.color.rgb = GRAY_MID

paragraph_space(doc, after=10)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 — IMPLEMENTATION
# ══════════════════════════════════════════════════════════════════════════════
banner_table(doc, '11  |  Implementation & Onboarding — Up and Running in 48 Hours')

add_body(doc,
    'We understand you are busy running a business.  Our onboarding is designed to get you '
    'productive with zero disruption to your daily operations.',
    size=11)

steps = [
    ('Day 1  —  Setup & Configuration (2–3 hours)',
     ['Store profile setup: branding, GST details, bank info',
      'Team creation: add staff with roles and login credentials',
      'Product catalog import: SKUs, categories, and opening stock',
      'Channel integrations: WhatsApp, Instagram, website chat',
      'Voice AI Agent configuration: language, scripts, and call flows']),
    ('Day 2  —  Training & Go-Live (3–4 hours)',
     ['Owner/Manager training: Dashboard, Leads, Billing, and Reports',
      'Staff training: Walk-ins, Clock-in/out, and Staff Portal',
      'First live walk-in captured and lead created in the system',
      'First billing transaction processed through POS',
      'Voice AI Agent tested with live inbound call']),
    ('Week 1  —  Stabilization & Optimization',
     ['Daily check-in support from onboarding team',
      'Lead pipeline populated from historical WhatsApp/notebook data',
      'Custom order templates set up for your standard project types',
      'First email campaign planned and scheduled',
      'Dashboard reviewed with owner — KPIs baseline established']),
    ('Ongoing  —  Growth Support',
     ['Monthly business review call with your dedicated account manager',
      'New feature updates rolled out automatically — no manual upgrades',
      'WhatsApp support channel for day-to-day queries',
      'Quarterly marketing strategy consultation']),
]
for title, points in steps:
    add_subheading(doc, title, color=BROWN_MID)
    for pt in points:
        add_bullet(doc, pt)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 12 — NEXT STEPS / CTA
# ══════════════════════════════════════════════════════════════════════════════
banner_table(doc, '12  |  Next Steps — Start Your Free Demo Today', bg_hex='C98B2B')

add_body(doc,
    'You have seen the problems.  You have seen the solution.  The only question is: '
    'how much longer can you afford to run your furniture business without it?',
    size=12, color=BROWN_DARK)

add_body(doc,
    'Every day without FurnitureCRM means:',
    size=11, color=DARK_TEXT)

cost_bullets = [
    '3–5 potential customers calling and getting no answer',
    'At least 1 follow-up forgotten — possibly a ₹50,000+ sale lost',
    'Billing taking 5x longer than it needs to',
    'Owner spending 3 hours on operations instead of growing the business',
    'Competitors using technology to move faster than you',
]
for cb in cost_bullets:
    add_pain_bullet(doc, cb)

paragraph_space(doc, after=10)
add_divider(doc, 'C98B2B')
paragraph_space(doc, after=10)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Book a Free 30-Minute Live Demo')
r.bold = True
r.font.size = Pt(20)
r.font.color.rgb = BROWN_DARK

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('We will walk you through the full system using your own store\'s scenario.\n'
              'No sales pressure. Just a clear look at what your business can achieve.')
r.font.size = Pt(11)
r.font.color.rgb = GRAY_MID
r.italic = True

paragraph_space(doc, after=12)

tbl = doc.add_table(rows=1, cols=3)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
contact_items = [
    ('📞', 'Call / WhatsApp', 'Your number here'),
    ('📧', 'Email',           'your@email.com'),
    ('🌐', 'Website',         'www.yourwebsite.com'),
]
for col_idx, (icon, label, value) in enumerate(contact_items):
    cell = tbl.cell(0, col_idx)
    set_cell_bg(cell, 'FDF6EC')
    add_cell_border(cell, 'C98B2B', 'bottom', '10')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(10)
    r  = p.add_run(icon + '\n')
    r.font.size = Pt(18)
    r2 = p.add_run(label + '\n')
    r2.bold = True
    r2.font.size = Pt(10)
    r2.font.color.rgb = BROWN_DARK
    r3 = p.add_run(value)
    r3.font.size = Pt(10)
    r3.font.color.rgb = GOLD

paragraph_space(doc, after=20)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('FurnitureCRM  —  Built for Indian Furniture Businesses  |  © 2025')
r.font.size = Pt(9)
r.font.color.rgb = GRAY_MID
r.italic = True

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r'c:\Users\divya\Desktop\Furzentic\FurnitureCRM_Business_Proposal.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
